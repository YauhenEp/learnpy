import datetime
from docx import Document
import os
from docx.opc.exceptions import PackageNotFoundError
import hw6
from hw8 import JsonWriter

class Publisher:
    def __init__(self):
        self.document_path = 'news_list.docx'

    options = ["News", "Private Add", "Fun story", "From File", "From Json"]

    def select_from_options(self, options):
        print("Please choose from the following options:")
        for i, option in enumerate(options, start=1):
            print(f"{i}. {option}")
        while True:
            try:
                choice = int(input("Enter the number of your choice: "))
                if 1 <= choice <= len(options):
                    selected_option = options[choice - 1]
                    print(f"You selected: {selected_option}")
                    break
                else:
                    print(f"Please enter a number between 1 and {len(options)}.")
            except ValueError:
                print("Invalid input. Please enter a number.")
        return selected_option

    def add_text_to_document(self, text):
        try:
            if not os.path.exists(self.document_path):
                raise FileNotFoundError(f"File not found: {self.document_path}")
            print(Document(self.document_path))
            template_document = Document(self.document_path)
            template_document.add_paragraph(text)
            template_document.save(self.document_path)
            print("Text added and document saved successfully.")
        except FileNotFoundError as e:
            print(f"Error: {e}")
        except PackageNotFoundError:
            print(f"Error: The file at {self.document_path} is not a valid .docx file.")
        except Exception as e:
            print(f"An unexpected error occurred: {e}")

    def start_publisher(self):
        item_to_add = ''
        selected_option = self.select_from_options(self.options)
        text = input(f'Input your {selected_option}: ')
        if selected_option == "News":
            location = input('Input your location: ')
            item = News(text, location)
            item_to_add = item.generate_news_text()
        if selected_option == "Private Add":
            item = PrivateAdd(text)
            item_to_add = item.generate_add_text()
        if selected_option == "Fun story":
            item = FunStory(text)
            item_to_add = item.generate_fun_story_text()
        if selected_option == "From File":
            item = hw6.From_File(text)
            item_to_add = item.generate_file_text()
        if selected_option == "From Json":
            json_type = input('1 From file; 2 From string')
            if(json_type == '1'):
                file_path = input('File Path - ')
                item = JsonWriter(file_path)
                item_to_add = item.generate_json()
            if(json_type == '2'):
                data_dict = input('Dict - ')
                item = JsonWriter(data_dict)
                item_to_add = item.generate_json()
        print('item_to_add ', item_to_add)
        self.add_text_to_document(item_to_add)


class News(Publisher):
    def __init__(self, news_text, location):
        self.news_text = news_text
        self.location = location

    def generate_news_text(self):
        news_text = f'''
News -------------------
{self.news_text}
Location: {self.location} {datetime.date.today()}
'''
        print(news_text)
        return news_text

class PrivateAdd(Publisher):
    def __init__(self, add_text):
        self.add_text = add_text

    def generate_add_text(self):
        news_text = f'''
Private Add -------------------
{self.add_text}
'''
        return news_text



class FunStory(Publisher):
    def __init__(self, funStory):
        self.funStory = funStory

    def generate_fun_story_text(self):
        story_text = f'''
Fun Story -------------------
{self.funStory}
{datetime.date.today()}
'''
        return story_text

publisher = Publisher()
publisher.start_publisher()